"""Import quiz questions from an uploaded PDF/Word document.

Two steps: extract plain text from the file (pypdf / python-docx), then ask Groq
(OpenAI-compatible) to return the questions as structured JSON. The result is a
*draft* the teacher reviews and edits before anything is saved.
"""
import io
import json
import logging

import requests

from app.config import GROQ_API_KEY, GROQ_MODEL

logger = logging.getLogger(__name__)

_GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
_TIMEOUT = 60
_MAX_CHARS = 24000  # keep well within the model context

VALID_TYPES = {"multiple_choice", "true_false", "short_answer"}


class QuizImportError(Exception):
    pass


def is_configured() -> bool:
    return bool(GROQ_API_KEY)


# --- text extraction --------------------------------------------------------
def extract_text(filename: str, content: bytes) -> str:
    ext = (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else ""
    if ext == "pdf":
        text = _pdf_text(content)
    elif ext in ("docx",):
        text = _docx_text(content)
    elif ext in ("txt", "md"):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise QuizImportError("Unsupported file type. Upload a PDF, Word (.docx), or text file.")
    text = text.strip()
    if not text:
        raise QuizImportError(
            "Couldn't read any text from that file. If it's a scanned image, "
            "export a text-based PDF or Word document instead."
        )
    return text[:_MAX_CHARS]


def _pdf_text(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx_text(content: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# --- AI extraction ----------------------------------------------------------
_SYSTEM = (
    "You extract quiz questions from a document's text and return STRICT JSON. "
    "Output an object with a single key \"questions\": an array. Each question is: "
    "{\"type\": one of \"multiple_choice\" | \"true_false\" | \"short_answer\", "
    "\"prompt\": string, \"points\": integer >= 1, "
    "\"options\": [{\"text\": string, \"is_correct\": boolean}], "
    "\"correct_answer\": string|null}. "
    "Rules: For multiple_choice include all answer choices as options and mark the "
    "correct one(s) is_correct=true. For true_false use options "
    "[{\"text\":\"True\",...},{\"text\":\"False\",...}]. For short_answer leave options "
    "empty and put the expected answer in correct_answer (or null if unknown). "
    "If the document doesn't mark the correct answer, make your best guess. "
    "Default points to 1. Return ONLY the JSON, no prose."
)


def extract_questions(text: str) -> list[dict]:
    if not is_configured():
        raise QuizImportError("AI import isn't configured on the server (GROQ_API_KEY missing).")
    try:
        resp = requests.post(
            _GROQ_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": GROQ_MODEL,
                "temperature": 0.1,
                "response_format": {"type": "json_object"},
                "messages": [
                    {"role": "system", "content": _SYSTEM},
                    {"role": "user", "content": f"Document text:\n\n{text}"},
                ],
            },
            timeout=_TIMEOUT,
        )
    except requests.RequestException as exc:
        raise QuizImportError(f"Couldn't reach the AI service: {exc}") from exc

    if resp.status_code != 200:
        logger.error("Groq error %s: %s", resp.status_code, resp.text[:500])
        raise QuizImportError("The AI service returned an error. Try again.")

    try:
        content = resp.json()["choices"][0]["message"]["content"]
        data = json.loads(content)
    except (KeyError, IndexError, ValueError) as exc:
        raise QuizImportError("The AI response couldn't be parsed.") from exc

    raw = data.get("questions") if isinstance(data, dict) else None
    if not isinstance(raw, list):
        raise QuizImportError("No questions were found in that document.")
    return [q for q in (_normalize(item) for item in raw) if q]


def _normalize(item: dict) -> dict | None:
    """Coerce one AI question into our QuestionInput shape; drop if unusable."""
    if not isinstance(item, dict):
        return None
    prompt = str(item.get("prompt") or "").strip()
    if not prompt:
        return None
    qtype = item.get("type")
    if qtype not in VALID_TYPES:
        qtype = "multiple_choice" if item.get("options") else "short_answer"

    try:
        points = max(1, int(item.get("points") or 1))
    except (TypeError, ValueError):
        points = 1

    options: list[dict] = []
    if qtype == "true_false":
        correct = item.get("correct_answer")
        tf_true = True
        if isinstance(correct, str):
            tf_true = correct.strip().lower() in ("true", "t", "yes")
        else:
            for o in item.get("options") or []:
                if isinstance(o, dict) and str(o.get("text", "")).strip().lower() == "true":
                    tf_true = bool(o.get("is_correct"))
        options = [
            {"text": "True", "is_correct": tf_true},
            {"text": "False", "is_correct": not tf_true},
        ]
    elif qtype == "multiple_choice":
        for o in item.get("options") or []:
            if isinstance(o, dict) and str(o.get("text") or "").strip():
                options.append({
                    "text": str(o["text"]).strip(),
                    "is_correct": bool(o.get("is_correct")),
                })
        if len(options) < 2:
            return None  # not a usable MC question
        if not any(o["is_correct"] for o in options):
            options[0]["is_correct"] = True  # ensure one correct

    correct_answer = item.get("correct_answer")
    correct_answer = str(correct_answer).strip() if correct_answer else None

    return {
        "type": qtype,
        "prompt": prompt,
        "points": points,
        "options": options,
        "correct_answer": correct_answer if qtype == "short_answer" else None,
    }
